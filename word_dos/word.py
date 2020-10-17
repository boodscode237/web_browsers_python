import docx
# import os
d = docx.Document('/home/abodo/Documents/webbrowser_python/word_dos/demo.docx')
# os.chdir('/home/abodo/Documents/webbrowser_python/word_dos/')
p=d.paragraphs[1]
# runs tells us everytime the styles changes
p.runs[2]
p.runs[3].underline = True
p.runs[3].text = 'underlined text'
d.save('/home/abodo/Documents/webbrowser_python/word_dos/demo2.docx')

# create a new doc
d = docx.Document()
d.add_paragraph('Hello my friend this is a new doc')
d.add_paragraph('Hello my bro this is a new new doc')
d.save('/home/abodo/Documents/webbrowser_python/word_dos/demo4.docx')

# function to return texts from doc in one line
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
    return '\n'.join(fullText)
